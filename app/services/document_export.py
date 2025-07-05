import os
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path
import json

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentExporter:
    def __init__(self, export_dir: str = "exports"):
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(exist_ok=True)
    
    def export_to_pdf(self, cover_letter_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """Export cover letter to PDF format"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is not installed. Install with: pip install reportlab")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"cover_letter_{timestamp}.pdf"
        
        filepath = self.export_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=12,
            leading=18
        )
        
        # Build content
        story = []
        
        # Title
        title = f"Cover Letter - {cover_letter_data.get('job_title', 'Position')} at {cover_letter_data.get('company_name', 'Company')}"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Date
        date_str = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(date_str, normal_style))
        story.append(Spacer(1, 20))
        
        # Content
        content = cover_letter_data.get('generated_content', '')
        if content:
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), normal_style))
                    story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        return str(filepath)
    
    def export_to_docx(self, cover_letter_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """Export cover letter to DOCX format"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install with: pip install python-docx")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"cover_letter_{timestamp}.docx"
        
        filepath = self.export_dir / filename
        
        # Create DOCX document
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(f"Cover Letter - {cover_letter_data.get('job_title', 'Position')} at {cover_letter_data.get('company_name', 'Company')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_str = datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(date_str)
        doc.add_paragraph()  # Empty paragraph for spacing
        
        # Add content
        content = cover_letter_data.get('generated_content', '')
        if content:
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # Save document
        doc.save(str(filepath))
        return str(filepath)
    
    def export_to_txt(self, cover_letter_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """Export cover letter to plain text format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"cover_letter_{timestamp}.txt"
        
        filepath = self.export_dir / filename
        
        # Build content
        lines = []
        lines.append(f"Cover Letter - {cover_letter_data.get('job_title', 'Position')} at {cover_letter_data.get('company_name', 'Company')}")
        lines.append("=" * 80)
        lines.append("")
        lines.append(datetime.now().strftime("%B %d, %Y"))
        lines.append("")
        
        content = cover_letter_data.get('generated_content', '')
        if content:
            lines.append(content)
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return str(filepath)
    
    def get_available_formats(self) -> Dict[str, bool]:
        """Get available export formats"""
        return {
            "pdf": REPORTLAB_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "txt": True  # Always available
        } 